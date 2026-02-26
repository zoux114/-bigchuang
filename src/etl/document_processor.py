"""
文档处理 ETL 模块
支持 PDF/Docx/TXT 转换为 Markdown，保留规章制度的层级结构
"""
import os
import re
import hashlib
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from datetime import datetime

# 文档处理库
try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from docx import Document
except ImportError:
    Document = None

from config.settings import (
    RAW_DATA_DIR,
    PROCESSED_DATA_DIR,
    ETL_CONFIG,
    CHUNKING_CONFIG,
)


class DocumentProcessor:
    """文档处理器 - 将原始文档转换为结构化 Markdown"""

    def __init__(self):
        self.supported_formats = ETL_CONFIG["supported_formats"]
        self.section_patterns = CHUNKING_CONFIG["section_patterns"]

    def process_all(self, force: bool = False) -> Dict[str, str]:
        """
        处理 raw 目录下的所有文档

        Args:
            force: 是否强制重新处理所有文件

        Returns:
            处理结果字典 {原文件名: 处理状态}
        """
        results = {}

        for file_path in RAW_DATA_DIR.iterdir():
            if file_path.suffix.lower() not in self.supported_formats:
                continue

            output_path = self._get_output_path(file_path)

            # 检查是否需要处理
            if not force and self._is_processed(file_path, output_path):
                results[file_path.name] = "skipped (already processed)"
                continue

            try:
                content = self._process_file(file_path)
                self._save_markdown(output_path, content, file_path.name)
                results[file_path.name] = "success"
            except Exception as e:
                results[file_path.name] = f"failed: {str(e)}"

        return results

    def process_single(self, file_path: Path) -> str:
        """处理单个文件并返回 Markdown 内容"""
        return self._process_file(file_path)

    def _process_file(self, file_path: Path) -> str:
        """根据文件类型选择处理方法"""
        suffix = file_path.suffix.lower()

        if suffix == ".pdf":
            return self._process_pdf(file_path)
        elif suffix == ".docx":
            return self._process_docx(file_path)
        elif suffix == ".txt":
            return self._process_txt(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {suffix}")

    def _process_pdf(self, file_path: Path) -> str:
        """处理 PDF 文件"""
        if pdfplumber is None:
            raise ImportError("请安装 pdfplumber: pip install pdfplumber")

        content_parts = []

        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                text = page.extract_text()
                if text:
                    # 清理文本
                    text = self._clean_text(text)
                    if text.strip():
                        content_parts.append(f"<!-- 第 {page_num} 页 -->\n{text}")

        return self._structure_content("\n\n".join(content_parts))

    def _process_docx(self, file_path: Path) -> str:
        """处理 Word 文档"""
        if Document is None:
            raise ImportError("请安装 python-docx: pip install python-docx")

        doc = Document(file_path)
        content_parts = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # 识别标题层级
                if para.style.name.startswith("Heading"):
                    level = int(para.style.name[-1]) if para.style.name[-1].isdigit() else 1
                    text = f"{'#' * level} {text}"
                content_parts.append(text)

        # 处理表格
        for table in doc.tables:
            table_md = self._table_to_markdown(table)
            if table_md:
                content_parts.append(table_md)

        return self._structure_content("\n\n".join(content_parts))

    def _process_txt(self, file_path: Path) -> str:
        """处理纯文本文件"""
        # 尝试多种编码
        encodings = ["utf-8", "gbk", "gb2312", "utf-16"]

        content = None
        for encoding in encodings:
            try:
                content = file_path.read_text(encoding=encoding)
                break
            except UnicodeDecodeError:
                continue

        if content is None:
            raise ValueError("无法识别文件编码")

        content = self._clean_text(content)
        return self._structure_content(content)

    def _clean_text(self, text: str) -> str:
        """清理文本内容"""
        # 移除多余的空白行
        text = re.sub(r"\n{3,}", "\n\n", text)
        # 移除行首行尾空白
        lines = [line.strip() for line in text.split("\n")]
        return "\n".join(lines)

    def _structure_content(self, content: str) -> str:
        """
        将文本内容结构化，识别规章制度的层级
        """
        lines = content.split("\n")
        structured_lines = []

        for line in lines:
            line = line.strip()
            if not line:
                structured_lines.append("")
                continue

            # 检测章节标题
            is_section = False
            for pattern in self.section_patterns:
                if re.match(pattern, line):
                    # 根据模式确定标题级别
                    if "章" in pattern:
                        line = f"## {line}"
                    elif "节" in pattern:
                        line = f"### {line}"
                    elif "条" in pattern:
                        line = f"#### {line}"
                    is_section = True
                    break

            # 如果不是章节，检测是否是其他标题特征
            if not is_section:
                # 数字编号的条款 (如 "1.1", "（一）")
                if re.match(r"^[\d]+\.[\d]+\s", line):
                    line = f"**{line}**"
                elif re.match(r"^[（(][一二三四五六七八九十]+[)）]", line):
                    line = f"- {line}"

            structured_lines.append(line)

        return "\n".join(structured_lines)

    def _table_to_markdown(self, table) -> str:
        """将 Word 表格转换为 Markdown 格式"""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")

        if not rows:
            return ""

        # 添加表头分隔符
        if len(rows) > 1:
            separator = "| " + " | ".join(["---"] * len(rows[0].split("|"))-2) + " |"
            rows.insert(1, separator)

        return "\n".join(rows)

    def _get_output_path(self, input_path: Path) -> Path:
        """获取输出文件路径"""
        return PROCESSED_DATA_DIR / f"{input_path.stem}.md"

    def _is_processed(self, input_path: Path, output_path: Path) -> bool:
        """检查文件是否已经处理过（基于修改时间）"""
        if not output_path.exists():
            return False

        input_mtime = input_path.stat().st_mtime
        output_mtime = output_path.stat().st_mtime

        return output_mtime > input_mtime

    def _save_markdown(self, output_path: Path, content: str, source_name: str):
        """保存 Markdown 文件"""
        # 添加元数据头部
        header = f"""---
source: {source_name}
processed_at: {datetime.now().isoformat()}
---

"""
        output_path.write_text(header + content, encoding="utf-8")


class DocumentWatcher:
    """文档监控器 - 监控 raw 目录变化"""

    def __init__(self, processor: DocumentProcessor):
        self.processor = processor

    def start_watching(self):
        """开始监控目录变化"""
        try:
            from watchdog.observers import Observer
            from watchdog.events import FileSystemEventHandler, FileCreatedEvent, FileModifiedEvent

            class Handler(FileSystemEventHandler):
                def __init__(self, processor):
                    self.processor = processor

                def on_created(self, event):
                    if not event.is_directory:
                        print(f"检测到新文件: {event.src_path}")
                        self._process_file(event.src_path)

                def on_modified(self, event):
                    if not event.is_directory:
                        print(f"检测到文件修改: {event.src_path}")
                        self._process_file(event.src_path)

                def _process_file(self, path):
                    file_path = Path(path)
                    if file_path.suffix.lower() in ETL_CONFIG["supported_formats"]:
                        try:
                            self.processor.process_single(file_path)
                            print(f"处理完成: {file_path.name}")
                        except Exception as e:
                            print(f"处理失败: {e}")

            observer = Observer()
            handler = Handler(self.processor)
            observer.schedule(handler, str(RAW_DATA_DIR), recursive=False)
            observer.start()

            print(f"开始监控目录: {RAW_DATA_DIR}")
            print("按 Ctrl+C 停止监控...")

            try:
                while True:
                    import time
                    time.sleep(1)
            except KeyboardInterrupt:
                observer.stop()
            observer.join()

        except ImportError:
            print("请安装 watchdog 以启用文件监控: pip install watchdog")


def main():
    """ETL 处理入口"""
    import argparse

    parser = argparse.ArgumentParser(description="文档处理 ETL")
    parser.add_argument("--force", action="store_true", help="强制重新处理所有文件")
    parser.add_argument("--watch", action="store_true", help="启用文件监控模式")
    args = parser.parse_args()

    processor = DocumentProcessor()

    if args.watch:
        watcher = DocumentWatcher(processor)
        watcher.start_watching()
    else:
        print("开始处理文档...")
        results = processor.process_all(force=args.force)
        print("\n处理结果:")
        for file_name, status in results.items():
            print(f"  {file_name}: {status}")


if __name__ == "__main__":
    main()
